#coding=utf-8
from docx import Document
import os
from constant import HOST,HOST_NEED_CHANGE
import re
from docx.shared import Inches



class WordTable():
    def __init__(self,table):
        self.table=table

    #获取功能名
    def get_name(self):
        return self.table.cell(0,1).text

    #获取请求头链接
    def get_url(self):
        checkurl=self.table.cell(1,1).text
        for key in HOST_NEED_CHANGE:
            if key in checkurl:
                checkurl=checkurl.replace(key,HOST)
        return checkurl

    #获取方法
    def get_method(self):
        zhmodel = re.compile(u'[\u4e00-\u9fa5]')#检查中文
        match = zhmodel.search(self.table.cell(2,4).text)
        if match:
            name=self.table.cell(2,4).text
        else:
            name=self.table.cell(2,5).text
        return name
    #获取表格信息
    def get_colums(self):
        table_list={}
        for i,row in enumerate(self.table.rows):
            row_content={}

            for cell in row.cells:
                text=cell.text
                row_content.append(text)
        table_list.append(row_content)

        return table_list
    #返回参数组合
    def get_params(self):
        start=end=0
        paramname=[u"请求协议",u"请求参数"]
        row_list=self.table.rows
        for i,row in enumerate(row_list):
            if self.table.cell(i,0).text in paramname:
                if start==end and end==0:
                    start=end=i
                elif start==end and end!=0:
                    end=i
                else:
                    end=i
        param_list=self.get_param(start+2,end+1)

        return param_list

    #返回chechkpoints数据
    def get_checkpoint(self):
        start = end = 0
        paramname = u"返回数据"
        first_list_name = []
        for i, row in enumerate(self.table.rows):
            if self.table.cell(i,0).text in paramname:
                if start == end and end == 0:
                    start = end = i
                elif start == end and end != 0:
                    end = i
                else:
                    end = i
        checkpoint_list = self.get_param(start, end+1)

        return checkpoint_list

    def _get_samplename(self):
        start = 0
        paramname = u"样例"
        first_list_name = []
        col_list = self.table.columns
        for i, row in enumerate(self.table.rows):
            if self.table.cell(i, 0).text in paramname:
                start=i
                break
        sample_list = self.get_sample(start)

        return sample_list
    #获取参数
    def get_param(self,start,end):
        param=[]

        for row in range(start,end):
            str=self.get_singlelineparam(row)
            param.append(str)
        return param

    #获取单行参数
    def get_singlelineparam(self,rownum):
        str = ""
        num = 6
        for i in range(1, num):

            if i == 2:
                str += ',#'
            elif i == 1:
                str += "\n'" + self.table.cell(rownum, i).text + "':"
            elif i==num:
                str += self.table.cell(rownum, i).text
            else:
                str += self.table.cell(rownum, i).text
        return str
    #获取返回数据
    def get_rdata(self,start,end):
        rdata=[]
        for row in range(start,end):
            str = self.get_singlelineparam(row)
            rdata.append(str)

        return rdata

    #获取样例
    def get_sample(self,rownum):
        return self.table.cell(rownum,1).text


data=["wordname","wordpath","tablenumber","content"]
testdata=["method","url","params","checkpoint","testcasename"]
class WordUtil():
    def __init__(self,path):
        self.path=path
        self.document=Document(path)
        self.tables = self.document.tables

    def get_tablesdata(self):
        num=0
        tables={}
        datas = []
        for table in self.tables:
            datas.append(self.get_testdata(table))
            num+=1
        tables[data[0]]=os.path.splitext(os.path.split(self.path)[1])[0]
        tables[data[1]]=self.path
        tables[data[2]]=num
        tables[data[3]]=datas

        return tables

    #获取单表信息
    def get_testdata(self,table):
        dict={}
        t = WordTable(table)
        dict[testdata[0]] = t.get_method()
        dict[testdata[1]] = t.get_url()
        dict[testdata[2]] = t.get_params()
        dict[testdata[3]] = t.get_checkpoint()
        #去掉样例
        #dict[testdata[4]] = t._get_samplename()
        dict[testdata[4]]=t.get_name()
        return dict
    # def close_word(self):
    #     self.document.Quit()


if __name__=="__main__":
    a=WordUtil("m2c.scm.api-1.3.0.docx")
    print a.get_tablesdata()

